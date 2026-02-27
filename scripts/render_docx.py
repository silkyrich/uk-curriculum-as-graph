#!/usr/bin/env python3
"""
DOCX renderer for teacher planners.

Generates formatted Word documents for any study node type.
Adapts layout detail by key stage:
  - KS1: Simpler tables, fewer sections, larger fonts
  - KS2: Full treatment matching the Roman Britain reference output
  - KS3-4: Denser layout, more analytical sections, smaller fonts

All content from graph — zero LLM.
"""

import datetime
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from planner_queries import StudyContext

# ── Subject colour palettes (hex strings for XML) ────────────────────

PALETTES = {
    'History':              {'primary': 'C0392B', 'accent': 'F39C12', 'bg': 'FDF2E9'},
    'Geography':            {'primary': '27AE60', 'accent': '2ECC71', 'bg': 'EAFAF1'},
    'Science':              {'primary': '2E86C1', 'accent': '3498DB', 'bg': 'EBF5FB'},
    'English':              {'primary': '8E44AD', 'accent': '9B59B6', 'bg': 'F4ECF7'},
    'Mathematics':          {'primary': 'E67E22', 'accent': 'F39C12', 'bg': 'FEF5E7'},
    'Art and Design':       {'primary': 'E74C3C', 'accent': 'F1948A', 'bg': 'FDEDEC'},
    'Music':                {'primary': '1ABC9C', 'accent': '48C9B0', 'bg': 'E8F8F5'},
    'Design and Technology': {'primary': 'D35400', 'accent': 'E67E22', 'bg': 'FDF2E9'},
    'Computing':            {'primary': '2C3E50', 'accent': '34495E', 'bg': 'EBEDEF'},
    'Religious Studies':    {'primary': '7D3C98', 'accent': 'AF7AC5', 'bg': 'F4ECF7'},
    'Citizenship':          {'primary': '16A085', 'accent': '1ABC9C', 'bg': 'E8F6F3'},
}

DARK_BROWN  = RGBColor(0x5D, 0x40, 0x37)
DARK_SLATE  = RGBColor(0x2C, 0x3E, 0x50)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
SKY_BLUE    = RGBColor(0x2E, 0x86, 0xC1)
PURPLE      = RGBColor(0x8E, 0x44, 0xAD)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
FOREST_GREEN = RGBColor(0x27, 0xAE, 0x60)

# ── Age-band parameters ──────────────────────────────────────────────

AGE_PARAMS = {
    'KS1': {'title_size': 28, 'heading_size': 14, 'body_size': 11, 'small_size': 9},
    'KS2': {'title_size': 32, 'heading_size': 13, 'body_size': 10, 'small_size': 9},
    'KS3': {'title_size': 28, 'heading_size': 12, 'body_size': 10, 'small_size': 8},
    'KS4': {'title_size': 28, 'heading_size': 12, 'body_size': 10, 'small_size': 8},
}

DIFF_LEVEL_COLOURS = {
    'entry': 'F39C12', 'developing': 'E67E22', 'expected': '27AE60',
    'greater_depth': '2E86C1', 'greater depth': '2E86C1',
    'emerging': 'F39C12', 'secure': '27AE60', 'mastery': '2E86C1',
}


def _hex_to_rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _get_pal(subject: str) -> dict:
    return PALETTES.get(subject, PALETTES['History'])


def _get_age(ks: str) -> dict:
    return AGE_PARAMS.get(ks, AGE_PARAMS['KS2'])


# ── DOCX helpers (from Roman Britain script) ─────────────────────────

def set_cell_shading(cell, colour_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colour_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_existing = tcPr.find(qn('w:tcMar'))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcPr.append(tcMar)


def add_formatted_para(doc_or_cell, text, size=11, colour=BLACK, bold=False,
                       italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=0, space_after=6, font_name='Calibri'):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = colour
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    return p


def add_mixed_para(doc_or_cell, parts, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before=0, space_after=6):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for text, size, colour, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = colour
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Calibri'
    return p


def add_coloured_heading(doc, text, level, colour):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = colour
        run.font.name = 'Calibri'
    return h


def add_horizontal_rule(doc, colour_hex='C0392B'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{colour_hex}"/>'
        f'</w:pBdr>'
    )
    pBdr_existing = pPr.find(qn('w:pBdr'))
    if pBdr_existing is not None:
        pPr.remove(pBdr_existing)
    pPr.append(pBdr)


def _as_list(val):
    """Ensure value is a list."""
    if val is None:
        return []
    if isinstance(val, str):
        try:
            parsed = json.loads(val)
            if isinstance(parsed, list):
                return parsed
        except (json.JSONDecodeError, TypeError):
            pass
        return [val] if val else []
    return val


# ── Section builders ─────────────────────────────────────────────────

def _build_title_block(doc, ctx, pal, age):
    """Title banner + metadata table."""
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_table.rows[0].cells[0]
    set_cell_shading(cell, pal['primary'])
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)

    add_formatted_para(cell, f'{ctx.subject.upper()}  |  TEACHER PLANNER', size=12,
                       colour=RGBColor(0xFF, 0xCC, 0xCC),
                       bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    study_name = ctx.study.get('name', '')
    add_formatted_para(cell, study_name, size=age['title_size'],
                       colour=WHITE, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
    if ctx.study_id:
        add_formatted_para(cell, f'[{ctx.study_id}]', size=11,
                           colour=RGBColor(0xFF, 0xCC, 0xCC),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)

    period = ctx.study.get('period', '')
    if period:
        add_formatted_para(cell, period, size=14, colour=_hex_to_rgb(pal['accent']),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)

    doc.add_paragraph()

    # Metadata table
    yg = ctx.study.get('year_groups', [])
    if isinstance(yg, list):
        yg = ', '.join(yg)
    duration = ctx.study.get('duration_lessons', '')
    study_type = ctx.study.get('study_type', ctx.study.get('unit_type', ctx.study.get('enquiry_type', '')))
    if study_type:
        study_type = study_type.replace('_', ' ').title()
    status = ctx.study.get('curriculum_status', '')
    cur_ref = ctx.study.get('curriculum_reference', [])
    if isinstance(cur_ref, list) and cur_ref:
        cur_ref = cur_ref[0]

    meta_data = [
        [('Subject', ctx.subject), ('Key Stage', ctx.key_stage),
         ('Year Group', yg or '-'), ('Duration', f"{duration} lessons" if duration else '-')],
        [('Study Type', study_type or '-'), ('Status', status.title() if status else '-'),
         ('NC Reference', str(cur_ref)[:50] if cur_ref else '-'), ('', '')],
    ]
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row_data in enumerate(meta_data):
        for col_idx, (label, value) in enumerate(row_data):
            if not label:
                continue
            cell = meta_table.rows[row_idx].cells[col_idx]
            set_cell_shading(cell, pal['bg'])
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            add_mixed_para(cell, [
                (label + ': ', 9, DARK_BROWN, True, False),
                (str(value), 10, DARK_SLATE, False, False),
            ], space_before=0, space_after=0)

    doc.add_paragraph()


def _build_enquiry_section(doc, ctx, pal, age):
    """Enquiry questions section."""
    eqs = _as_list(ctx.study.get('enquiry_questions'))
    if not eqs:
        eq = ctx.study.get('enquiry_question', '')
        if eq:
            eqs = [eq]
    if not eqs:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Enquiry Questions', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    for i, q in enumerate(eqs, 1):
        add_mixed_para(doc, [
            (f'{i}.  ', 12, _hex_to_rgb(pal['accent']), True, False),
            (q, 12, DARK_SLATE, False, True),
        ], space_after=4)


def _build_concepts_section(doc, ctx, pal, age):
    """Concepts with differentiation tables."""
    if not ctx.concepts:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Concepts', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    primary = [c for c in ctx.concepts if c.get('is_primary')]
    secondary = [c for c in ctx.concepts if not c.get('is_primary')]

    # Full treatment for primary concepts
    for c in primary:
        _build_concept_full(doc, c, pal, age, is_primary=True)

    # Compact for secondary
    if secondary:
        add_coloured_heading(doc, 'Supporting Concepts', level=3, colour=DARK_SLATE)
        for c in secondary:
            if len(secondary) <= 5:
                _build_concept_compact(doc, c, pal, age)
            else:
                _build_concept_minimal(doc, c, pal, age)

    doc.add_paragraph()


def _build_concept_full(doc, concept, pal, age, is_primary=True):
    """Full concept rendering with teaching guidance + differentiation table."""
    name = concept.get('concept_name', concept.get('name', ''))
    cid = concept.get('concept_id', '')
    ctype = concept.get('concept_type', '')
    tw = concept.get('teaching_weight', '')

    label = 'Primary' if is_primary else 'Secondary'
    add_coloured_heading(doc, f'{name}', level=3, colour=DARK_SLATE)
    add_mixed_para(doc, [
        (cid, 9, SKY_BLUE, False, False),
        (f'  |  {label} concept', 9, DARK_BROWN, False, False),
        (f'  |  {ctype.title()}' if ctype else '', 9, DARK_BROWN, False, False),
        (f'  |  Weight: {tw}/6' if tw else '', 9, DARK_BROWN, False, False),
    ], space_after=8)

    desc = concept.get('description', '')
    if desc:
        add_formatted_para(doc, desc, size=age['body_size'], colour=DARK_SLATE)

    # Teaching guidance box
    guidance = concept.get('teaching_guidance', '')
    if guidance:
        tg_table = doc.add_table(rows=1, cols=1)
        cell = tg_table.rows[0].cells[0]
        set_cell_shading(cell, 'EBF5FB')
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        add_formatted_para(cell, 'TEACHING GUIDANCE', size=9, colour=SKY_BLUE, bold=True, space_after=4)
        add_formatted_para(cell, guidance, size=age['body_size'], colour=DARK_SLATE)

    # Misconceptions box
    misconceptions = concept.get('common_misconceptions', '')
    if misconceptions:
        mc_table = doc.add_table(rows=1, cols=1)
        cell = mc_table.rows[0].cells[0]
        set_cell_shading(cell, 'FDEDEC')
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        add_formatted_para(cell, 'COMMON MISCONCEPTIONS', size=9,
                           colour=_hex_to_rgb(pal['primary']), bold=True, space_after=4)
        add_formatted_para(cell, misconceptions, size=age['body_size'], colour=DARK_SLATE)

    # Key vocabulary box
    kv = concept.get('key_vocabulary', '')
    if kv:
        kv_table = doc.add_table(rows=1, cols=1)
        cell = kv_table.rows[0].cells[0]
        set_cell_shading(cell, 'F5EEF8')
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        add_formatted_para(cell, 'KEY VOCABULARY', size=9, colour=PURPLE, bold=True, space_after=4)
        add_formatted_para(cell, kv if isinstance(kv, str) else ', '.join(kv),
                           size=age['body_size'], colour=DARK_SLATE, italic=True)

    # Differentiation table
    dls = concept.get('difficulty_levels', [])
    if dls:
        doc.add_paragraph()
        add_mixed_para(doc, [
            ('Differentiation  ', 11, DARK_SLATE, True, False),
            ('(from DifficultyLevel nodes)', 9, DARK_BROWN, False, True),
        ], space_after=4)

        has_task = any(dl.get('example_task') for dl in dls)
        has_response = any(dl.get('example_response') for dl in dls)
        has_errors = any(dl.get('common_errors') for dl in dls)

        cols = 2 + (1 if has_task else 0) + (1 if has_errors else 0)
        diff_table = doc.add_table(rows=len(dls) + 1, cols=cols)
        diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ['Level', 'What success looks like']
        if has_task:
            headers.append('Example task')
        if has_errors:
            headers.append('Common errors')

        for i, header in enumerate(headers):
            cell = diff_table.rows[0].cells[i]
            set_cell_shading(cell, '2C3E50')
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            add_formatted_para(cell, header, size=9, colour=WHITE, bold=True, space_after=0)

        # Data rows
        for row_idx, dl in enumerate(dls):
            row = diff_table.rows[row_idx + 1]
            label = dl.get('label', '').replace('_', ' ')
            level_colour = DIFF_LEVEL_COLOURS.get(label.lower(), pal['primary'])

            # Level cell
            cell = row.cells[0]
            set_cell_shading(cell, level_colour)
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            add_formatted_para(cell, label.title(), size=9, colour=WHITE, bold=True,
                               space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Description cell
            cell = row.cells[1]
            bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            add_formatted_para(cell, dl.get('description', ''), size=9,
                               colour=DARK_SLATE, space_after=0)
            # Model response sub-row within the description cell
            ex_resp = dl.get('example_response', '')
            if ex_resp and has_response:
                add_formatted_para(cell, f'Model response: {ex_resp}', size=8,
                                   colour=DARK_BROWN, italic=True, space_before=4, space_after=0)

            col_offset = 2
            if has_task:
                cell = row.cells[col_offset]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                add_formatted_para(cell, dl.get('example_task', ''), size=9,
                                   colour=DARK_SLATE, space_after=0)
                col_offset += 1

            if has_errors:
                cell = row.cells[col_offset]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                errors = dl.get('common_errors', '')
                if isinstance(errors, list):
                    errors = '; '.join(errors)
                add_formatted_para(cell, errors, size=9, colour=DARK_SLATE, space_after=0)

    # CPA representation stages
    rss = concept.get('representation_stages', [])
    if rss:
        doc.add_paragraph()
        add_mixed_para(doc, [
            ('CPA Stages  ', 11, DARK_SLATE, True, False),
            ('(Concrete-Pictorial-Abstract)', 9, DARK_BROWN, False, True),
        ], space_after=4)

        rs_table = doc.add_table(rows=len(rss) + 1, cols=4)
        rs_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(['Stage', 'Description', 'Resources', 'Transition cue']):
            cell = rs_table.rows[0].cells[i]
            set_cell_shading(cell, '06B6D4')
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            add_formatted_para(cell, h, size=9, colour=WHITE, bold=True, space_after=0)

        for row_idx, rs in enumerate(rss):
            row = rs_table.rows[row_idx + 1]
            bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'

            cell = row.cells[0]
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            add_formatted_para(cell, rs.get('stage', '').title(), size=9,
                               colour=DARK_SLATE, bold=True, space_after=0)

            for col_idx, key in enumerate(['description', 'resources', 'transition_cue'], 1):
                cell = row.cells[col_idx]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
                val = rs.get(key, '')
                if isinstance(val, list):
                    val = ', '.join(val)
                add_formatted_para(cell, val, size=9, colour=DARK_SLATE, space_after=0)

    doc.add_paragraph()


def _build_concept_compact(doc, concept, pal, age):
    """Compact concept — name, ID, description, no table."""
    name = concept.get('concept_name', concept.get('name', ''))
    cid = concept.get('concept_id', '')
    add_mixed_para(doc, [
        (f'{name}  ', 11, DARK_SLATE, True, False),
        (f'({cid})', 9, SKY_BLUE, False, False),
    ], space_after=2)
    desc = concept.get('description', '')
    if desc:
        add_formatted_para(doc, desc[:200] + ('...' if len(desc) > 200 else ''),
                           size=age['body_size'], colour=DARK_BROWN, italic=True, space_after=8)


def _build_concept_minimal(doc, concept, pal, age):
    """One-line concept."""
    name = concept.get('concept_name', concept.get('name', ''))
    cid = concept.get('concept_id', '')
    desc = concept.get('description', '')
    add_mixed_para(doc, [
        (f'{name} ', 10, DARK_SLATE, True, False),
        (f'({cid}) ', 8, SKY_BLUE, False, False),
        (desc[:80] + '...' if len(desc) > 80 else desc, 9, DARK_BROWN, False, True),
    ], space_after=4)


def _build_thinking_lens_section(doc, ctx, pal, age):
    """Thinking lens section."""
    if not ctx.thinking_lenses:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Thinking Lens', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    primary_lens = ctx.thinking_lenses[0]
    lens_table = doc.add_table(rows=1, cols=1)
    cell = lens_table.rows[0].cells[0]
    set_cell_shading(cell, 'F5EEF8')
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    add_mixed_para(cell, [
        ('PRIMARY:  ', 10, PURPLE, True, False),
        (primary_lens.get('lens_name', ''), 12, PURPLE, True, False),
    ], space_after=6)
    add_formatted_para(cell, primary_lens.get('key_question', ''),
                       size=11, colour=DARK_SLATE, italic=True, space_after=8)

    if primary_lens.get('rationale'):
        add_formatted_para(cell, primary_lens['rationale'],
                           size=age['body_size'], colour=DARK_BROWN, space_after=8)

    stems = primary_lens.get('question_stems', [])
    if stems:
        add_formatted_para(cell, f'Question stems for {ctx.key_stage}:',
                           size=9, colour=PURPLE, bold=True, space_after=4)
        for s in stems:
            add_formatted_para(cell, f'    •  {s}', size=age['body_size'],
                               colour=DARK_SLATE, space_after=2)

    # Secondary lens
    if len(ctx.thinking_lenses) > 1:
        doc.add_paragraph()
        sec = ctx.thinking_lenses[1]
        add_mixed_para(doc, [
            ('Secondary lens:  ', 10, DARK_BROWN, True, False),
            (sec['lens_name'], 10, PURPLE, True, False),
            (f" — {sec.get('rationale', '')}" if sec.get('rationale') else '', 10, DARK_BROWN, False, False),
        ])

    doc.add_paragraph()


def _build_session_structure(doc, ctx, pal, age):
    """Vehicle template / session structure section."""
    if not ctx.templates:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Session Structure', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    for vt in ctx.templates:
        name = vt.get('name', 'Template')
        phases = vt.get('session_structure', [])
        assessment = vt.get('assessment_approach', '')

        colour_map = {
            'topic_study': FOREST_GREEN, 'source_enquiry': SKY_BLUE,
            'fair_test': RGBColor(0x2E, 0x86, 0xC1), 'text_study': PURPLE,
        }
        vt_colour = colour_map.get(vt.get('template_type', ''), _hex_to_rgb(pal['accent']))

        add_mixed_para(doc, [(name, 12, vt_colour, True, False)], space_after=4)

        if phases:
            add_formatted_para(doc, '  →  '.join(phases), size=age['body_size'],
                               colour=vt_colour, bold=True, space_after=4)

        if assessment:
            add_mixed_para(doc, [
                ('Assessment: ', 9, DARK_BROWN, True, False),
                (assessment, 9, DARK_BROWN, False, True),
            ], space_after=10)

    # Question stems from templates
    for vt in ctx.templates:
        stems = vt.get('ks_question_stems', [])
        if stems:
            add_formatted_para(doc, f'{ctx.key_stage} Question Stems:',
                               size=age['body_size'], colour=DARK_SLATE, bold=True, space_after=4)
            for q in stems:
                add_formatted_para(doc, f'    •  {q}', size=age['body_size'],
                                   colour=DARK_SLATE, space_after=2)
            break

    doc.add_paragraph()


def _build_subject_specific(doc, ctx, pal, age):
    """Subject-specific sections."""
    primary_colour = _hex_to_rgb(pal['primary'])
    study = ctx.study

    if ctx.subject == 'History':
        _build_history_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Science':
        _build_science_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'English':
        _build_english_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Geography':
        _build_geography_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Art and Design':
        _build_art_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Music':
        _build_music_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Design and Technology':
        _build_dt_sections(doc, ctx, pal, age, primary_colour)
    elif ctx.subject == 'Computing':
        _build_computing_sections(doc, ctx, pal, age, primary_colour)


def _build_history_sections(doc, ctx, pal, age, primary_colour):
    """History-specific: sources, figures, disciplinary concepts."""
    study = ctx.study

    # Primary sources
    sources = ctx.references.get('sources', [])
    if sources:
        add_coloured_heading(doc, 'Primary Sources', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])

        source_colours = ['EBF5FB', pal['bg'], 'F2F3F4']
        source_accents = [SKY_BLUE, _hex_to_rgb(pal['accent']), DARK_SLATE]

        for i, src in enumerate(sources):
            src_table = doc.add_table(rows=1, cols=1)
            cell = src_table.rows[0].cells[0]
            set_cell_shading(cell, source_colours[i % len(source_colours)])
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

            name = src.get('name', src.get('source_name', ''))
            stype = src.get('source_type', '').replace('_', ' ').title()
            date = src.get('date', src.get('date_range', ''))

            add_mixed_para(cell, [
                (f'{name}  ', 13, source_accents[i % len(source_accents)], True, False),
                (f'({stype}, {date})' if stype else '', 10, DARK_BROWN, False, True),
            ], space_after=6)

            prov = src.get('provenance', src.get('description', ''))
            if prov:
                add_formatted_para(cell, prov, size=age['body_size'], colour=DARK_SLATE, space_after=8)

            ped = src.get('pedagogical_use', '')
            if ped:
                add_formatted_para(cell, 'How to use in the classroom:', size=9,
                                   colour=source_accents[i % len(source_accents)], bold=True, space_after=4)
                add_formatted_para(cell, ped, size=age['body_size'], colour=DARK_BROWN, space_after=6)

            loc = src.get('location', '')
            if loc:
                add_mixed_para(cell, [('Location: ', 9, DARK_BROWN, True, False),
                                      (loc, 9, DARK_BROWN, False, False)], space_after=2)
            url = src.get('url', '')
            if url:
                add_mixed_para(cell, [('URL: ', 9, DARK_BROWN, True, False),
                                      (url, 9, SKY_BLUE, False, False)], space_after=0)

            doc.add_paragraph()

    # Key figures and events
    figures = _as_list(study.get('key_figures'))
    events = _as_list(study.get('key_events'))
    if figures or events:
        add_coloured_heading(doc, 'Key Figures and Events', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])

        if figures:
            fig_table = doc.add_table(rows=len(figures) + 1, cols=2)
            fig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(['Figure', 'Significance']):
                cell = fig_table.rows[0].cells[i]
                set_cell_shading(cell, '2C3E50')
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                add_formatted_para(cell, h, size=10, colour=WHITE, bold=True, space_after=0)
            for row_idx, fig in enumerate(figures):
                row = fig_table.rows[row_idx + 1]
                bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'
                cell = row.cells[0]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
                add_formatted_para(cell, fig, size=10, colour=DARK_SLATE, bold=True, space_after=0)
                cell = row.cells[1]
                set_cell_shading(cell, bg)
                set_cell_margins(cell, top=50, bottom=50, left=80, right=80)

            doc.add_paragraph()

        # Perspectives
        perspectives = _as_list(study.get('perspectives'))
        if perspectives:
            add_mixed_para(doc, [
                ('Perspectives to include: ', 10, DARK_SLATE, True, False),
                ('  •  '.join(perspectives), 10, DARK_BROWN, False, True),
            ])

        # Significance claim
        sig = study.get('significance_claim', '')
        if sig:
            doc.add_paragraph()
            sig_table = doc.add_table(rows=1, cols=1)
            cell = sig_table.rows[0].cells[0]
            set_cell_shading(cell, '2C3E50')
            set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
            add_formatted_para(cell, 'SIGNIFICANCE', size=10,
                               colour=_hex_to_rgb(pal['accent']), bold=True, space_after=4)
            add_formatted_para(cell, sig, size=11, colour=WHITE, italic=True)

        doc.add_paragraph()


def _build_science_sections(doc, ctx, pal, age, primary_colour):
    """Science-specific: variables, equipment, safety."""
    study = ctx.study

    variables = study.get('variables')
    if isinstance(variables, str):
        try:
            variables = json.loads(variables)
        except (json.JSONDecodeError, TypeError):
            variables = None

    if variables and isinstance(variables, dict):
        add_coloured_heading(doc, 'Variables', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        add_mixed_para(doc, [('Independent: ', 10, primary_colour, True, False),
                             (variables.get('independent', ''), 10, DARK_SLATE, False, False)])
        add_mixed_para(doc, [('Dependent: ', 10, primary_colour, True, False),
                             (variables.get('dependent', ''), 10, DARK_SLATE, False, False)])
        controlled = variables.get('controlled', [])
        if controlled:
            if isinstance(controlled, list):
                controlled = ', '.join(controlled)
            add_mixed_para(doc, [('Controlled: ', 10, primary_colour, True, False),
                                 (controlled, 10, DARK_SLATE, False, False)])
        doc.add_paragraph()

    equipment = _as_list(study.get('equipment'))
    safety = study.get('safety_notes', '')
    if equipment or safety:
        add_coloured_heading(doc, 'Equipment and Safety', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        for e in equipment:
            add_formatted_para(doc, f'    •  {e}', size=age['body_size'], colour=DARK_SLATE, space_after=2)
        if safety:
            doc.add_paragraph()
            sc_table = doc.add_table(rows=1, cols=1)
            cell = sc_table.rows[0].cells[0]
            set_cell_shading(cell, 'FDF2E9')
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            add_formatted_para(cell, 'SAFETY', size=9, colour=ORANGE, bold=True, space_after=4)
            add_formatted_para(cell, safety, size=age['body_size'], colour=DARK_BROWN)
        doc.add_paragraph()


def _build_english_sections(doc, ctx, pal, age, primary_colour):
    """English-specific: text type, grammar, suggested texts."""
    study = ctx.study

    text_type = study.get('text_type', '')
    features = _as_list(study.get('text_features_to_teach'))
    grammar = _as_list(study.get('grammar_focus'))
    outcome = study.get('writing_outcome', '')
    texts = study.get('suggested_texts', [])
    if isinstance(texts, str):
        try:
            texts = json.loads(texts)
        except (json.JSONDecodeError, TypeError):
            texts = []

    if text_type or features or grammar or outcome:
        add_coloured_heading(doc, 'Text Study Focus', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        if text_type:
            add_mixed_para(doc, [('Text type: ', 10, DARK_SLATE, True, False),
                                 (text_type.replace('_', ' ').title(), 10, DARK_BROWN, False, False)])
        if outcome:
            add_mixed_para(doc, [('Writing outcome: ', 10, DARK_SLATE, True, False),
                                 (outcome, 10, DARK_BROWN, False, False)])
        if grammar:
            add_mixed_para(doc, [('Grammar focus: ', 10, DARK_SLATE, True, False),
                                 (', '.join(grammar), 10, DARK_BROWN, False, False)])
        if features:
            add_mixed_para(doc, [('Features to teach: ', 10, DARK_SLATE, True, False),
                                 (', '.join(features), 10, DARK_BROWN, False, False)])
        doc.add_paragraph()

    if texts:
        add_coloured_heading(doc, 'Suggested Texts', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        for t in texts:
            if isinstance(t, dict):
                add_mixed_para(doc, [
                    (t.get('title', ''), 10, DARK_SLATE, True, True),
                    (f" by {t.get('author', '')}", 10, DARK_BROWN, False, False),
                    (f" — {t.get('note', '')}" if t.get('note') else '', 9, DARK_BROWN, False, True),
                ], space_after=4)
        doc.add_paragraph()


def _build_geography_sections(doc, ctx, pal, age, primary_colour):
    """Geography: scale, maps, fieldwork."""
    study = ctx.study
    scale = study.get('scale', '')
    themes = _as_list(study.get('themes'))
    maps = _as_list(study.get('map_types'))
    fieldwork = study.get('fieldwork_potential', '')

    if scale or themes or maps or fieldwork:
        add_coloured_heading(doc, 'Study Scope', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        if scale:
            add_mixed_para(doc, [('Scale: ', 10, DARK_SLATE, True, False),
                                 (scale.replace('_', ' ').title(), 10, DARK_BROWN, False, False)])
        if themes:
            add_mixed_para(doc, [('Themes: ', 10, DARK_SLATE, True, False),
                                 (', '.join(themes), 10, DARK_BROWN, False, False)])
        if maps:
            add_mixed_para(doc, [('Map types: ', 10, DARK_SLATE, True, False),
                                 (', '.join(m.replace('_', ' ') for m in maps), 10, DARK_BROWN, False, False)])
        if fieldwork:
            add_mixed_para(doc, [('Fieldwork: ', 10, DARK_SLATE, True, False),
                                 (fieldwork, 10, DARK_BROWN, False, False)])
        doc.add_paragraph()


def _build_art_sections(doc, ctx, pal, age, primary_colour):
    """Art: artist, medium, techniques."""
    study = ctx.study
    add_coloured_heading(doc, 'Art Focus', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])
    fields = [
        ('Artist', study.get('artist', '')),
        ('Art movement', study.get('art_movement', '')),
        ('Medium', ', '.join(_as_list(study.get('medium')))),
        ('Techniques', ', '.join(_as_list(study.get('techniques')))),
        ('Visual elements', ', '.join(_as_list(study.get('visual_elements')))),
        ('Cultural context', study.get('cultural_context', '')),
    ]
    for label, val in fields:
        if val:
            add_mixed_para(doc, [(f'{label}: ', 10, DARK_SLATE, True, False),
                                 (val, 10, DARK_BROWN, False, False)])
    doc.add_paragraph()


def _build_music_sections(doc, ctx, pal, age, primary_colour):
    """Music: genre, composer, elements."""
    study = ctx.study
    add_coloured_heading(doc, 'Music Focus', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])
    fields = [
        ('Genre', (study.get('genre') or '').replace('_', ' ').title()),
        ('Composer/piece', f"{study.get('composer', '')}" +
         (f" — {study.get('piece', '')}" if study.get('piece') and study.get('piece') != study.get('composer') else '')),
        ('Musical elements', ', '.join(_as_list(study.get('musical_elements')))),
        ('Instruments', ', '.join(_as_list(study.get('instrument')))),
        ('Notation level', (study.get('notation_level') or '').replace('_', ' ')),
        ('MMC reference', study.get('mmc_reference', '')),
    ]
    for label, val in fields:
        if val and val.strip():
            add_mixed_para(doc, [(f'{label}: ', 10, DARK_SLATE, True, False),
                                 (val, 10, DARK_BROWN, False, False)])
    doc.add_paragraph()


def _build_dt_sections(doc, ctx, pal, age, primary_colour):
    """DT: design brief, materials, tools."""
    study = ctx.study
    strand = (study.get('dt_strand') or '').replace('_', ' ').title()
    add_coloured_heading(doc, f'Design and Technology: {strand}' if strand else 'Design and Technology',
                         level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])
    brief = study.get('design_brief', '')
    if brief:
        add_formatted_para(doc, brief, size=age['body_size'], colour=DARK_SLATE)
    fields = [
        ('Materials', ', '.join(_as_list(study.get('materials')))),
        ('Tools', ', '.join(_as_list(study.get('tools')))),
        ('Techniques', ', '.join(_as_list(study.get('techniques')))),
        ('Safety', study.get('safety_notes', '')),
    ]
    for label, val in fields:
        if val:
            add_mixed_para(doc, [(f'{label}: ', 10, DARK_SLATE, True, False),
                                 (val, 10, DARK_BROWN, False, False)])
    # Food-specific
    allergens = _as_list(study.get('food_allergens'))
    if allergens:
        add_mixed_para(doc, [('Allergens: ', 10, ORANGE, True, False),
                             (', '.join(allergens), 10, DARK_BROWN, False, False)])
    doc.add_paragraph()


def _build_computing_sections(doc, ctx, pal, age, primary_colour):
    """Computing: paradigm, tool, concepts."""
    study = ctx.study
    add_coloured_heading(doc, 'Computing Focus', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])
    fields = [
        ('Programming paradigm', (study.get('programming_paradigm') or '').replace('_', ' ').title()),
        ('Software/tool', study.get('software_tool', '')),
        ('Computational concepts', ', '.join(c.replace('_', ' ') for c in _as_list(study.get('computational_concept')))),
        ('Abstraction level', (study.get('abstraction_level') or '').replace('_', ' ').title()),
    ]
    for label, val in fields:
        if val and val.strip():
            add_mixed_para(doc, [(f'{label}: ', 10, DARK_SLATE, True, False),
                                 (val, 10, DARK_BROWN, False, False)])
    doc.add_paragraph()


def _build_sequencing(doc, ctx, pal, age):
    """Sequencing section — before/this/after."""
    if not ctx.follows and not ctx.leads_to:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Sequencing', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    seq_table = doc.add_table(rows=1, cols=3)
    seq_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    items = [
        ('BEFORE', ctx.follows or '—', 'F2F3F4', DARK_SLATE),
        ('THIS STUDY', ctx.study.get('name', ''), pal['primary'], WHITE),
        ('AFTER', ctx.leads_to or '—', 'F2F3F4', DARK_SLATE),
    ]
    for i, (label, study_name, bg, text_col) in enumerate(items):
        cell = seq_table.rows[0].cells[i]
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        label_col = _hex_to_rgb(pal['accent']) if i == 1 else text_col
        add_formatted_para(cell, label, size=9, colour=label_col, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        add_formatted_para(cell, study_name, size=11, colour=text_col, bold=(i == 1),
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_paragraph()


def _build_pitfalls_and_sensitive(doc, ctx, pal, age):
    """Pitfalls and sensitive content."""
    pitfalls = _as_list(ctx.study.get('common_pitfalls'))
    sensitive = _as_list(ctx.study.get('sensitive_content_notes'))
    primary_colour = _hex_to_rgb(pal['primary'])

    if pitfalls:
        add_coloured_heading(doc, 'Pitfalls to Avoid', level=2, colour=primary_colour)
        add_horizontal_rule(doc, pal['primary'])
        for p in pitfalls:
            add_formatted_para(doc, f'    \u2717  {p}', size=age['body_size'],
                               colour=primary_colour, space_after=4)
        doc.add_paragraph()

    if sensitive:
        sc_table = doc.add_table(rows=1, cols=1)
        cell = sc_table.rows[0].cells[0]
        set_cell_shading(cell, 'FDF2E9')
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        add_formatted_para(cell, 'SENSITIVE CONTENT', size=9, colour=ORANGE, bold=True, space_after=6)
        for s in sensitive:
            add_formatted_para(cell, f'•  {s}', size=age['body_size'], colour=DARK_BROWN, space_after=2)
        doc.add_paragraph()


def _build_cross_curricular(doc, ctx, pal, age):
    """Cross-curricular opportunities table."""
    if not ctx.cross_curricular:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Cross-Curricular Opportunities', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    xc_table = doc.add_table(rows=len(ctx.cross_curricular) + 1, cols=3)
    xc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Subject', 'Link', 'Connection']):
        cell = xc_table.rows[0].cells[i]
        set_cell_shading(cell, '2C3E50')
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        add_formatted_para(cell, h, size=10, colour=WHITE, bold=True, space_after=0)

    for row_idx, cc in enumerate(ctx.cross_curricular):
        row = xc_table.rows[row_idx + 1]
        bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'
        data = [
            cc.get('target_subject', cc.get('target_label', '')),
            cc.get('target_name', ''),
            cc.get('hook', ''),
        ]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            bold = (col_idx == 0)
            add_formatted_para(cell, str(text), size=10, colour=DARK_SLATE,
                               bold=bold, space_after=0)

    doc.add_paragraph()


def _build_prior_knowledge(doc, ctx, pal, age):
    """Prior knowledge section — prerequisites grouped by target concept."""
    prereqs = getattr(ctx, 'prerequisites', [])
    if not prereqs:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Prior Knowledge', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    add_formatted_para(doc, 'Prerequisite knowledge students need before this study.',
                       size=9, colour=DARK_BROWN, italic=True, space_after=8)

    # Group by target concept
    grouped = {}
    for p in prereqs:
        target = p.get('target_name', 'Unknown')
        grouped.setdefault(target, []).append(p)

    pk_table = doc.add_table(rows=len(prereqs) + 1, cols=3)
    pk_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Prior knowledge needed', 'For concept', 'Description']):
        cell = pk_table.rows[0].cells[i]
        set_cell_shading(cell, '2C3E50')
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        add_formatted_para(cell, h, size=9, colour=WHITE, bold=True, space_after=0)

    row_idx = 0
    for target_name, items in grouped.items():
        for p in items:
            row = pk_table.rows[row_idx + 1]
            bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'

            cell = row.cells[0]
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            add_formatted_para(cell, p.get('prereq_name', ''), size=9,
                               colour=DARK_SLATE, bold=True, space_after=0)

            cell = row.cells[1]
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            add_formatted_para(cell, target_name, size=9,
                               colour=DARK_BROWN, space_after=0)

            cell = row.cells[2]
            set_cell_shading(cell, bg)
            set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
            add_formatted_para(cell, p.get('prereq_description', ''), size=9,
                               colour=DARK_SLATE, space_after=0)

            row_idx += 1

    doc.add_paragraph()


def _build_assessment_alignment(doc, ctx, pal, age):
    """Assessment alignment section — KS2 ContentDomainCodes only."""
    codes = getattr(ctx, 'assessment_codes', [])
    if not codes:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Assessment Alignment (KS2)', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    add_formatted_para(doc, 'KS2 test framework content domain codes covered by this study.',
                       size=9, colour=DARK_BROWN, italic=True, space_after=8)

    ac_table = doc.add_table(rows=len(codes) + 1, cols=3)
    ac_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Code', 'Description', 'Assesses concept']):
        cell = ac_table.rows[0].cells[i]
        set_cell_shading(cell, '2C3E50')
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        add_formatted_para(cell, h, size=9, colour=WHITE, bold=True, space_after=0)

    for row_idx, code in enumerate(codes):
        row = ac_table.rows[row_idx + 1]
        bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'

        cell = row.cells[0]
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        add_formatted_para(cell, code.get('code_id', ''), size=9,
                           colour=SKY_BLUE, bold=True, space_after=0)

        cell = row.cells[1]
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        add_formatted_para(cell, code.get('description', code.get('name', '')), size=9,
                           colour=DARK_SLATE, space_after=0)

        cell = row.cells[2]
        set_cell_shading(cell, bg)
        set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
        add_formatted_para(cell, code.get('concept_name', ''), size=9,
                           colour=DARK_BROWN, space_after=0)

    doc.add_paragraph()


def _build_scaffolding_inclusion(doc, ctx, pal, age):
    """Scaffolding and inclusion section — from learner profile data."""
    lp = getattr(ctx, 'learner_profile', {})
    if not lp:
        return

    pedagogy = lp.get('pedagogy', {})
    content = lp.get('content', {})
    feedback = lp.get('feedback', {})
    if not pedagogy and not content and not feedback:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Scaffolding and Inclusion', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    year_id = lp.get('year_id', '')
    if year_id:
        add_formatted_para(doc, f'Guidelines for {year_id}',
                           size=9, colour=DARK_BROWN, italic=True, space_after=8)

    # Build guideline rows
    rows_data = []

    # Reading level
    if content:
        label = content.get('label', '')
        lexile_min = content.get('lexile_min', '')
        lexile_max = content.get('lexile_max', '')
        lexile_range = f'{lexile_min}L–{lexile_max}L' if lexile_min and lexile_max else ''
        tts = content.get('tts_available', False)
        tts_note = content.get('tts_notes', '')
        reading_detail = label
        if lexile_range:
            reading_detail += f'  ({lexile_range})'
        if tts:
            reading_detail += f'  |  TTS available'
            if tts_note:
                reading_detail += f': {tts_note}'
        if reading_detail:
            rows_data.append(('Reading level', reading_detail))

        vocab = content.get('vocabulary_level', '')
        vocab_notes = content.get('vocabulary_notes', '')
        if vocab:
            vocab_detail = vocab
            if vocab_notes:
                vocab_detail += f' — {vocab_notes}'
            rows_data.append(('Vocabulary', vocab_detail))

        sentence_max = content.get('max_sentence_length_words', '')
        if sentence_max:
            rows_data.append(('Max sentence length', f'{sentence_max} words'))

    # Scaffolding
    if pedagogy:
        scaff = pedagogy.get('scaffolding_level', '')
        hints = pedagogy.get('hint_tiers_max', '')
        hint_notes = pedagogy.get('hint_tier_notes', '')
        if scaff:
            scaff_detail = str(scaff)
            if hints:
                scaff_detail += f'  |  Up to {hints} hint tiers'
            if hint_notes:
                scaff_detail += f'  ({hint_notes})'
            rows_data.append(('Scaffolding level', scaff_detail))

        session_min = pedagogy.get('session_length_min_minutes', '')
        session_max = pedagogy.get('session_length_max_minutes', '')
        if session_min and session_max:
            rows_data.append(('Session length', f'{session_min}–{session_max} minutes'))

        worked = pedagogy.get('worked_examples_required', False)
        if worked:
            style = pedagogy.get('worked_example_style', '')
            rows_data.append(('Worked examples', f'Required ({style})' if style else 'Required'))

        pf = pedagogy.get('productive_failure_appropriate', False)
        if pf:
            pf_notes = pedagogy.get('productive_failure_notes', '')
            rows_data.append(('Productive failure', pf_notes if pf_notes else 'Appropriate'))

    # Feedback
    if feedback:
        tone = feedback.get('ai_tone', '')
        normalize = feedback.get('normalize_struggle', False)
        error_framing = feedback.get('positive_error_framing', False)
        fb_detail_parts = []
        if tone:
            fb_detail_parts.append(tone)
        if normalize:
            fb_detail_parts.append('normalise struggle')
        if error_framing:
            fb_detail_parts.append('positive error framing')
        if fb_detail_parts:
            rows_data.append(('Feedback approach', '  |  '.join(fb_detail_parts)))

        ex_correct = feedback.get('feedback_example_correct', '')
        if ex_correct:
            rows_data.append(('Example (correct)', ex_correct))

        ex_incorrect = feedback.get('feedback_example_incorrect', '')
        if ex_incorrect:
            rows_data.append(('Example (incorrect)', ex_incorrect))

    if not rows_data:
        return

    sc_table = doc.add_table(rows=len(rows_data), cols=2)
    sc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, (guideline, detail) in enumerate(rows_data):
        bg = 'FDFEFE' if row_idx % 2 == 0 else 'F8F9F9'

        cell_l = sc_table.rows[row_idx].cells[0]
        set_cell_shading(cell_l, bg)
        set_cell_margins(cell_l, top=50, bottom=50, left=80, right=40)
        add_formatted_para(cell_l, guideline, size=9, colour=DARK_SLATE,
                           bold=True, space_after=0)

        cell_r = sc_table.rows[row_idx].cells[1]
        set_cell_shading(cell_r, bg)
        set_cell_margins(cell_r, top=50, bottom=50, left=40, right=80)
        add_formatted_para(cell_r, str(detail), size=9, colour=DARK_BROWN, space_after=0)

    doc.add_paragraph()


def _build_knowledge_organiser(doc, ctx, pal, age):
    """Knowledge organiser — compact assembly of key terms, events, figures, period, core facts."""
    study = ctx.study

    definitions = _as_list(study.get('definitions'))
    key_events = _as_list(study.get('key_events'))
    key_figures = _as_list(study.get('key_figures'))
    period = study.get('period', '')

    # Core facts: 'expected' or 'secure' difficulty level descriptions from primary concepts
    core_facts = []
    for c in ctx.concepts:
        if not c.get('is_primary'):
            continue
        dls = c.get('difficulty_levels', [])
        for dl in dls:
            lbl = (dl.get('label', '') or '').lower().replace('_', ' ')
            if lbl in ('expected', 'secure'):
                desc = dl.get('description', '')
                if desc:
                    core_facts.append((c.get('concept_name', c.get('name', '')), desc))
                break

    # Skip entirely if nothing to show
    if not definitions and not key_events and not key_figures and not period and not core_facts:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Knowledge Organiser', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    # Period
    if period:
        add_mixed_para(doc, [
            ('Period: ', 10, DARK_SLATE, True, False),
            (period, 10, DARK_BROWN, False, False),
        ], space_after=8)

    # Key terms — compact grid
    if definitions:
        add_formatted_para(doc, 'Key Terms', size=10, colour=DARK_SLATE,
                           bold=True, space_after=4)
        terms = []
        for item in definitions:
            if isinstance(item, dict):
                terms.append(item.get('term', item.get('word', '')))
            else:
                terms.append(str(item))
        terms = [t for t in terms if t]

        if terms:
            cols = min(4, len(terms))
            rows = (len(terms) + cols - 1) // cols
            term_table = doc.add_table(rows=rows, cols=cols)
            term_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, term in enumerate(terms):
                r, c = i // cols, i % cols
                cell = term_table.rows[r].cells[c]
                set_cell_shading(cell, pal['bg'])
                set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
                add_formatted_para(cell, term, size=9, colour=_hex_to_rgb(pal['primary']),
                                   bold=True, space_after=0,
                                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()

    # Key events — bulleted list
    if key_events:
        add_formatted_para(doc, 'Timeline / Key Events', size=10, colour=DARK_SLATE,
                           bold=True, space_after=4)
        for evt in key_events:
            add_formatted_para(doc, f'    \u2022  {evt}', size=age['body_size'],
                               colour=DARK_BROWN, space_after=2)
        doc.add_paragraph()

    # Key figures — comma-separated
    if key_figures:
        add_mixed_para(doc, [
            ('Key Figures: ', 10, DARK_SLATE, True, False),
            (', '.join(str(f) for f in key_figures), 10, DARK_BROWN, False, False),
        ], space_after=8)

    # Core facts — expected-level descriptions for primary concepts
    if core_facts:
        add_formatted_para(doc, 'Core Facts (Expected Standard)', size=10,
                           colour=DARK_SLATE, bold=True, space_after=4)
        for concept_name, desc in core_facts:
            add_mixed_para(doc, [
                (f'{concept_name}: ', 9, _hex_to_rgb(pal['primary']), True, False),
                (desc, 9, DARK_BROWN, False, False),
            ], space_after=4)

    doc.add_paragraph()


def _build_vocabulary(doc, ctx, pal, age):
    """Vocabulary word mat."""
    definitions = _as_list(ctx.study.get('definitions'))
    if not definitions:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Vocabulary Word Mat', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])
    add_formatted_para(doc, 'Print and place on desks. Introduce 3-4 words per lesson.',
                       size=9, colour=DARK_BROWN, italic=True, space_after=8)

    cols = 3
    rows = (len(definitions) + cols - 1) // cols
    vocab_table = doc.add_table(rows=rows, cols=cols)
    vocab_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    accent_colours = [pal['primary'], '8E44AD', '2E86C1', 'E67E22', '27AE60', pal['accent']]

    for i, item in enumerate(definitions):
        row_idx = i // cols
        col_idx = i % cols
        cell = vocab_table.rows[row_idx].cells[col_idx]
        set_cell_shading(cell, 'FDFEFE')
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)

        if isinstance(item, dict):
            word = item.get('term', item.get('word', ''))
            defn = item.get('meaning', item.get('definition', ''))
        else:
            word = str(item)
            defn = ''

        accent = accent_colours[i % len(accent_colours)]
        word_colour = _hex_to_rgb(accent) if isinstance(accent, str) else accent
        add_formatted_para(cell, word, size=11, colour=word_colour, bold=True, space_after=2)
        if defn:
            add_formatted_para(cell, defn, size=9, colour=DARK_SLATE, space_after=0)

    doc.add_paragraph()


def _build_success_criteria(doc, ctx, pal, age):
    """Success criteria section."""
    success = _as_list(ctx.study.get('success_criteria'))
    if not success:
        return

    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Success Criteria', level=2, colour=primary_colour)
    add_horizontal_rule(doc, pal['primary'])

    for s in success:
        add_formatted_para(doc, f'    \u2713  {s}', size=age['body_size'],
                           colour=FOREST_GREEN, space_after=4)

    doc.add_paragraph()


def _build_graph_context(doc, ctx, pal, age):
    """Graph context section — node IDs, Cypher queries, and data lookup hints for teachers."""
    primary_colour = _hex_to_rgb(pal['primary'])
    add_coloured_heading(doc, 'Graph Context (Data Lookup)', level=2, colour=DARK_SLATE)
    add_horizontal_rule(doc, '7F8C8D')

    add_formatted_para(doc, 'Use these identifiers to query the curriculum knowledge graph for additional data.',
                       size=age['small_size'], colour=DARK_BROWN, italic=True, space_after=8)

    # Node info table
    id_field = {
        'HistoryStudy': 'study_id', 'GeoStudy': 'study_id',
        'ScienceEnquiry': 'enquiry_id', 'EnglishUnit': 'unit_id',
    }.get(ctx.label, 'suggestion_id')

    domain_ids = ctx.study.get('domain_ids', [])
    if isinstance(domain_ids, str):
        try:
            domain_ids = json.loads(domain_ids)
        except (json.JSONDecodeError, TypeError):
            domain_ids = []
    template_ids = ctx.study.get('uses_template', [])
    if isinstance(template_ids, str):
        template_ids = [template_ids]

    info_rows = [
        ('Node type', ctx.label),
        ('Study ID', ctx.study_id),
        ('Subject', ctx.subject),
        ('Key stage', ctx.key_stage),
        ('Concepts delivered', f"{len(ctx.concepts)} ({sum(1 for c in ctx.concepts if c.get('is_primary'))} primary)"),
    ]
    if domain_ids:
        info_rows.append(('Domain IDs', ', '.join(domain_ids[:6])))
    if template_ids:
        info_rows.append(('Template IDs', ', '.join(template_ids)))
    if ctx.thinking_lenses:
        info_rows.append(('Primary thinking lens', ctx.thinking_lenses[0].get('lens_name', '')))
    if ctx.epistemic_skills:
        skill_names = [s.get('name', '') for s in ctx.epistemic_skills[:4]]
        info_rows.append(('Disciplinary skills', ', '.join(skill_names)))

    info_table = doc.add_table(rows=len(info_rows), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_rows):
        cell_l = info_table.rows[i].cells[0]
        cell_r = info_table.rows[i].cells[1]
        set_cell_margins(cell_l, top=30, bottom=30, left=80, right=40)
        set_cell_margins(cell_r, top=30, bottom=30, left=40, right=80)
        if i % 2 == 0:
            set_cell_shading(cell_l, 'F8F9F9')
            set_cell_shading(cell_r, 'F8F9F9')
        add_formatted_para(cell_l, label, size=age['small_size'],
                           colour=DARK_SLATE, bold=True, space_after=0, space_before=0)
        add_formatted_para(cell_r, value, size=age['small_size'],
                           colour=DARK_BROWN, space_after=0, space_before=0)

    doc.add_paragraph()

    # Concept IDs
    if ctx.concepts:
        add_formatted_para(doc, 'Concept IDs:', size=age['small_size'],
                           colour=DARK_SLATE, bold=True, space_after=4)
        for c in ctx.concepts:
            cid = c.get('concept_id', '')
            name = c.get('concept_name', c.get('name', ''))
            pri = ' (PRIMARY)' if c.get('is_primary') else ''
            add_formatted_para(doc, f'  {cid}: {name}{pri}',
                               size=age['small_size'], colour=DARK_BROWN, space_after=1)

    doc.add_paragraph()

    # Cypher query hints
    add_formatted_para(doc, 'Cypher query hints:', size=age['small_size'],
                       colour=DARK_SLATE, bold=True, space_after=4)

    query_box = doc.add_table(rows=1, cols=1)
    cell = query_box.rows[0].cells[0]
    set_cell_shading(cell, 'F8F9F9')
    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    queries = [
        f"// Find this study node",
        f"MATCH (ts:{ctx.label} {{{id_field}: '{ctx.study_id}'}})",
        f"RETURN ts",
        f"",
        f"// Get concepts + difficulty levels",
        f"MATCH (ts:{ctx.label} {{{id_field}: '{ctx.study_id}'}})",
        f"  -[:DELIVERS_VIA]->(c:Concept)",
        f"  -[:HAS_DIFFICULTY_LEVEL]->(dl)",
        f"RETURN c.name, dl.label, dl.description",
    ]
    for line in queries:
        add_formatted_para(cell, line, size=8, colour=DARK_SLATE,
                           font_name='Courier New', space_after=0, space_before=0)

    doc.add_paragraph()


def _build_footer(doc, ctx, pal):
    """Footer with provenance."""
    add_horizontal_rule(doc, pal['primary'])

    footer_table = doc.add_table(rows=1, cols=1)
    cell = footer_table.rows[0].cells[0]
    set_cell_shading(cell, 'F8F9F9')
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    add_formatted_para(cell, 'Generated from the UK Curriculum Knowledge Graph. All content sourced from Neo4j graph data — zero LLM generation.',
                       size=8, colour=DARK_BROWN, italic=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_formatted_para(cell, f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}',
                       size=8, colour=DARK_BROWN, italic=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    # Source document reference
    if ctx.source_documents:
        sd = ctx.source_documents[0]
        ref = sd.get('name', sd.get('document_name', ''))
        doc_ref = sd.get('reference', sd.get('document_reference', ''))
        if ref:
            add_formatted_para(cell, f'Source document: {ref}' + (f' ({doc_ref})' if doc_ref else ''),
                               size=8, colour=DARK_BROWN,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


# ── Main render function ─────────────────────────────────────────────

def render_docx(ctx: StudyContext) -> Document:
    """Render a complete DOCX document for a study. Returns Document object."""
    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    pal = _get_pal(ctx.subject)
    age = _get_age(ctx.key_stage)

    # Build sections
    _build_title_block(doc, ctx, pal, age)
    _build_enquiry_section(doc, ctx, pal, age)
    _build_concepts_section(doc, ctx, pal, age)
    _build_thinking_lens_section(doc, ctx, pal, age)
    _build_session_structure(doc, ctx, pal, age)
    _build_subject_specific(doc, ctx, pal, age)
    _build_sequencing(doc, ctx, pal, age)
    _build_pitfalls_and_sensitive(doc, ctx, pal, age)
    _build_cross_curricular(doc, ctx, pal, age)
    _build_prior_knowledge(doc, ctx, pal, age)
    _build_assessment_alignment(doc, ctx, pal, age)
    _build_scaffolding_inclusion(doc, ctx, pal, age)
    _build_knowledge_organiser(doc, ctx, pal, age)
    _build_vocabulary(doc, ctx, pal, age)
    _build_success_criteria(doc, ctx, pal, age)
    _build_graph_context(doc, ctx, pal, age)
    _build_footer(doc, ctx, pal)

    return doc
